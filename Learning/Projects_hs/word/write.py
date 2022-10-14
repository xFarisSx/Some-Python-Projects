import docx
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

mydoc = docx.Document()
mydoc.add_paragraph("This is the first")
mydoc.add_paragraph("This is the second")
mydoc.add_heading("FFff")
mydoc.add_picture('image.png', width=docx.shared.Inches(5),height=docx.shared.Inches(7))

mydoc.save('files/write.docx')