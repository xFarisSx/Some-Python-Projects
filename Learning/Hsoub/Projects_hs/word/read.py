import docx
from pathlib import Path
import readD

doc = docx.Document(Path('files', 'academy_1.docx'))

print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[2].text)

print(len(doc.paragraphs[2].runs))
print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)
print(doc.paragraphs[2].runs[4].text)
print('*'*50)
def getText(filename):
    doc = docx.Document(Path('files', filename))
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(readD.getText('academy_1.docx'))